"""Scientific Knowledge Platform: real structural extraction (title/abstract/headings/
references/tables), section-level search with confidence, document grounding, and the
Mission Knowledge Agent's citation-grounded (and honestly-refusing) answers."""

import fitz
import pytest


def _build_structured_pdf() -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Angular Rate Safety Verification via Farkas Certificates", fontsize=18)
    page.insert_text((72, 110), "Abstract", fontsize=14)
    page.insert_text((72, 130), "This paper presents a deterministic verifier for angular rate safety bounds using Farkas lemma.", fontsize=10)
    page.insert_text((72, 160), "1. Introduction", fontsize=14)
    page.insert_text((72, 180), "Spacecraft safety verification requires deterministic guarantees.", fontsize=10)
    page2 = doc.new_page()
    page2.insert_text((72, 72), "References", fontsize=14)
    page2.insert_text((72, 100), "[1] Farkas, J. Theorie der einfachen Ungleichungen. 1902.", fontsize=10)
    page2.insert_text((72, 120), "[2] Boyd, S. Convex Optimization. 2004.", fontsize=10)
    raw = doc.tobytes()
    doc.close()
    return raw


@pytest.fixture
def mission(client):
    return client.post("/api/missions", json={
        "mission_name": "Knowledge Platform Target", "mission_profile_key": "earth_observation",
    }).json()


def test_pdf_structure_extraction_is_real_not_fabricated(client, mission):
    pdf_bytes = _build_structured_pdf()
    resp = client.post(
        f"/api/missions/{mission['id']}/documents", params={"doc_type": "RESEARCH_PAPER"},
        files={"file": ("paper.pdf", pdf_bytes, "application/pdf")},
    )
    body = resp.json()
    summary = body["structure_summary"]
    assert summary["title"] == "Angular Rate Safety Verification via Farkas Certificates"
    assert summary["has_abstract"] is True
    assert summary["heading_count"] == 4
    assert summary["reference_count"] == 2
    assert summary["equations_note"] == "Equation extraction unavailable."


def test_document_sections_are_persisted_with_real_page_numbers(client, mission):
    pdf_bytes = _build_structured_pdf()
    upload = client.post(
        f"/api/missions/{mission['id']}/documents", params={"doc_type": "RESEARCH_PAPER"},
        files={"file": ("paper.pdf", pdf_bytes, "application/pdf")},
    ).json()

    sections = client.get(f"/api/missions/{mission['id']}/documents/{upload['document_id']}/sections").json()
    types = {s["section_type"] for s in sections}
    assert types == {"TITLE", "ABSTRACT", "HEADING", "REFERENCE"}

    refs = [s for s in sections if s["section_type"] == "REFERENCE"]
    assert len(refs) == 2
    assert "Farkas, J." in refs[0]["content_text"]

    references_heading = next(s for s in sections if s["section_type"] == "HEADING" and s["content_text"] == "References")
    assert references_heading["page_number"] == 1  # second page, 0-indexed


def test_docx_structure_uses_real_heading_styles(client, mission):
    from docx import Document
    import io
    doc = Document()
    doc.add_heading("Reaction Wheel Desaturation Algorithm", level=1)
    doc.add_paragraph("Abstract")
    doc.add_paragraph("This document describes a desaturation strategy for reaction wheels.")
    doc.add_heading("References", level=1)
    doc.add_paragraph("[1] Wie, B. Space Vehicle Dynamics and Control. 1998.")
    buf = io.BytesIO()
    doc.save(buf)

    resp = client.post(
        f"/api/missions/{mission['id']}/documents", params={"doc_type": "ALGORITHM_DESCRIPTION"},
        files={"file": ("algo.docx", buf.getvalue(),
                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    body = resp.json()
    assert body["extraction_status"] == "OK"
    assert body["structure_summary"]["heading_count"] == 2


def test_markdown_headings_extracted_exactly(client, mission):
    md = "# Mission Overview\n\nSome text.\n\n## Constraints\n\nMore text.\n"
    resp = client.post(f"/api/missions/{mission['id']}/documents", params={"doc_type": "ENGINEERING_NOTES"},
                        files={"file": ("notes.md", md, "text/markdown")})
    body = resp.json()
    assert body["structure_summary"]["heading_count"] == 2


def test_latex_sections_extracted_exactly(client, mission):
    tex = r"\title{Orbit Propagation Notes}\section{Introduction}Text.\subsection{SGP4}More text."
    resp = client.post(f"/api/missions/{mission['id']}/documents", params={"doc_type": "ALGORITHM_DESCRIPTION"},
                        files={"file": ("notes.tex", tex, "text/plain")})
    body = resp.json()
    assert body["structure_summary"]["title"] == "Orbit Propagation Notes"
    assert body["structure_summary"]["heading_count"] == 2


def test_corrupt_pdf_never_fabricates_structure(client, mission):
    resp = client.post(f"/api/missions/{mission['id']}/documents", params={"doc_type": "RESEARCH_PAPER"},
                        files={"file": ("broken.pdf", b"not a real pdf", "application/pdf")})
    body = resp.json()
    assert body["extraction_status"] == "CORRUPT"
    assert body["structure_summary"]["title"] is None
    assert body["structure_summary"]["heading_count"] == 0


def test_section_search_returns_document_page_section_confidence(client, mission):
    pdf_bytes = _build_structured_pdf()
    client.post(f"/api/missions/{mission['id']}/documents", params={"doc_type": "RESEARCH_PAPER"},
                files={"file": ("paper.pdf", pdf_bytes, "application/pdf")})

    resp = client.get(f"/api/missions/{mission['id']}/documents/sections/search", params={"q": "Farkas lemma"})
    body = resp.json()
    assert body["search_type"] == "keyword"
    assert len(body["results"]) >= 1
    top = body["results"][0]
    assert top["filename"] == "paper.pdf"
    assert "section_type" in top
    assert "page_number" in top
    assert 0 <= top["confidence"] <= 1.0


def test_section_search_filtered_by_type(client, mission):
    pdf_bytes = _build_structured_pdf()
    client.post(f"/api/missions/{mission['id']}/documents", params={"doc_type": "RESEARCH_PAPER"},
                files={"file": ("paper.pdf", pdf_bytes, "application/pdf")})

    resp = client.get(f"/api/missions/{mission['id']}/documents/sections/search",
                       params={"q": "Farkas", "section_type": "REFERENCE"})
    results = resp.json()["results"]
    assert all(r["section_type"] == "REFERENCE" for r in results)


def test_knowledge_agent_refuses_without_evidence(client, mission):
    resp = client.post(f"/api/missions/{mission['id']}/knowledge/ask", json={"question": "What is quantum gravity?"})
    body = resp.json()
    assert body["generated_by"] == "deterministic_refusal"
    assert body["answer"] == "I cannot find this information in the uploaded documents."
    assert body["evidence"] == []


def test_knowledge_agent_refuses_unrelated_question_even_with_documents_present(client, mission):
    """Regression: stopwords ('is', 'the', 'of', 'what') previously inflated match scores
    on ANY uploaded text, defeating the refusal for genuinely unrelated questions —
    caught by live-testing against a real server, not anticipated in advance."""
    pdf_bytes = _build_structured_pdf()
    client.post(f"/api/missions/{mission['id']}/documents", params={"doc_type": "RESEARCH_PAPER"},
                files={"file": ("paper.pdf", pdf_bytes, "application/pdf")})

    resp = client.post(f"/api/missions/{mission['id']}/knowledge/ask",
                        json={"question": "What is the capital of France?"})
    body = resp.json()
    assert body["generated_by"] == "deterministic_refusal"
    assert body["evidence"] == []


def test_knowledge_agent_finds_real_evidence_when_present(client, mission):
    """ANTHROPIC_API_KEY is cleared for the whole test suite (eval/conftest.py), so this
    exercises the real 'evidence found, Claude unavailable' fallback path — never a mock."""
    pdf_bytes = _build_structured_pdf()
    client.post(f"/api/missions/{mission['id']}/documents", params={"doc_type": "RESEARCH_PAPER"},
                files={"file": ("paper.pdf", pdf_bytes, "application/pdf")})

    resp = client.post(f"/api/missions/{mission['id']}/knowledge/ask",
                        json={"question": "What does the paper say about Farkas lemma?"})
    body = resp.json()
    assert body["generated_by"] == "deterministic_fallback"
    assert len(body["evidence"]) >= 1
    assert body["evidence"][0]["filename"] == "paper.pdf"


def test_knowledge_agent_empty_question_rejected(client, mission):
    resp = client.post(f"/api/missions/{mission['id']}/knowledge/ask", json={"question": "   "})
    assert resp.status_code == 400


def test_knowledge_agent_run_persisted_to_audit_trail(client, mission):
    client.post(f"/api/missions/{mission['id']}/knowledge/ask", json={"question": "anything"})
    runs = client.get(f"/api/missions/{mission['id']}/agent-runs",
                       params={"agent_name": "mission_knowledge_agent"}).json()
    assert len(runs) == 1
    assert runs[0]["status"] == "OK"


def test_knowledge_agent_missing_mission_404s(client):
    resp = client.post("/api/missions/999999/knowledge/ask", json={"question": "x"})
    assert resp.status_code == 404
