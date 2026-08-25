"""Real text + structure extraction for Mission Files. Every branch either extracts
genuine content or returns an honest status (UNSUPPORTED / CORRUPT / NOT_APPLICABLE) —
never a placeholder string standing in for extracted_text. Extraction is based on file
extension/content, not the operator-chosen doc_type label."""

import json
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from backend.documents.structure import (
    extract_docx_structure, extract_latex_structure, extract_markdown_structure, extract_pdf_structure,
)

TEXT_EXTENSIONS = {".txt", ".py", ".m", ".json", ".yaml", ".yml", ".ini", ".cfg", ".toml", ".csv"}
MARKDOWN_EXTENSIONS = {".md"}
LATEX_EXTENSIONS = {".tex"}
NOTEBOOK_EXTENSIONS = {".ipynb"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp"}

OK = "OK"
UNSUPPORTED = "UNSUPPORTED"
CORRUPT = "CORRUPT"
NOT_APPLICABLE = "NOT_APPLICABLE"

_EMPTY_STRUCTURE = {"headings": [], "tables": [], "references": [], "title": None, "abstract": None,
                     "equations_note": None, "reference_section_found": False}


def _extract_pdf(raw: bytes) -> tuple[str, str | None, dict, dict]:
    try:
        doc = fitz.open(stream=raw, filetype="pdf")
        pages_text = [page.get_text() for page in doc]
        metadata = {"page_count": doc.page_count, "pdf_metadata": {k: v for k, v in doc.metadata.items() if v}}
        structure = extract_pdf_structure(doc)
        doc.close()
        return OK, "\n\n".join(pages_text), metadata, structure
    except Exception as exc:  # PyMuPDF raises various exception types for malformed PDFs
        return CORRUPT, None, {"error": str(exc)}, _EMPTY_STRUCTURE


def _extract_docx(raw: bytes) -> tuple[str, str | None, dict, dict]:
    import io
    try:
        document = DocxDocument(io.BytesIO(raw))
        text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        metadata = {"paragraph_count": len(document.paragraphs), "table_count": len(document.tables)}
        structure = extract_docx_structure(document)
        return OK, text, metadata, structure
    except Exception as exc:
        return CORRUPT, None, {"error": str(exc)}, _EMPTY_STRUCTURE


def _extract_notebook(raw: bytes) -> tuple[str, str | None, dict, dict]:
    try:
        nb = json.loads(raw.decode("utf-8"))
        cells = nb.get("cells", [])
        parts = []
        for cell in cells:
            source = cell.get("source", [])
            text = "".join(source) if isinstance(source, list) else str(source)
            parts.append(f"[{cell.get('cell_type', 'unknown')}]\n{text}")
        metadata = {"cell_count": len(cells), "kernel": nb.get("metadata", {}).get("kernelspec", {}).get("name")}
        return OK, "\n\n".join(parts), metadata, _EMPTY_STRUCTURE
    except (json.JSONDecodeError, UnicodeDecodeError, AttributeError) as exc:
        return CORRUPT, None, {"error": str(exc)}, _EMPTY_STRUCTURE


def _extract_plain_text(raw: bytes, ext: str) -> tuple[str, str | None, dict, dict]:
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError as exc:
        return CORRUPT, None, {"error": f"not valid UTF-8 text: {exc}"}, _EMPTY_STRUCTURE

    metadata = {"char_count": len(text), "line_count": text.count("\n") + 1}
    if ext in MARKDOWN_EXTENSIONS:
        structure = extract_markdown_structure(text)
    elif ext in LATEX_EXTENSIONS:
        structure = extract_latex_structure(text)
    else:
        structure = _EMPTY_STRUCTURE
    return OK, text, metadata, structure


def extract(filename: str, raw: bytes) -> dict:
    """Returns {'status', 'text', 'metadata', 'structure'}. status is one of
    OK/UNSUPPORTED/CORRUPT/NOT_APPLICABLE; text is None unless status == OK. 'structure'
    (headings/tables/references/title/abstract/equations_note) is only populated for
    formats where real structural extraction is implemented (PDF, DOCX, Markdown, LaTeX)."""
    ext = Path(filename).suffix.lower()

    if ext in PDF_EXTENSIONS:
        status, text, metadata, structure = _extract_pdf(raw)
    elif ext in DOCX_EXTENSIONS:
        status, text, metadata, structure = _extract_docx(raw)
    elif ext in NOTEBOOK_EXTENSIONS:
        status, text, metadata, structure = _extract_notebook(raw)
    elif ext in TEXT_EXTENSIONS or ext in MARKDOWN_EXTENSIONS or ext in LATEX_EXTENSIONS:
        status, text, metadata, structure = _extract_plain_text(raw, ext)
    elif ext in IMAGE_EXTENSIONS:
        status, text, metadata, structure = (NOT_APPLICABLE, None,
                                              {"reason": "images require OCR/vision extraction, not implemented"}, _EMPTY_STRUCTURE)
    else:
        status, text, metadata, structure = UNSUPPORTED, None, {"extension": ext or "(none)"}, _EMPTY_STRUCTURE

    return {"status": status, "text": text, "metadata": metadata, "structure": structure}
